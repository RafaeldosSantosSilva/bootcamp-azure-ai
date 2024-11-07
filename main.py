import requests
from docx import Document
import os
from bs4 import BeautifulSoup
from langchain_openai.chat_models.azure import AzureChatOpenAI





# Fazendo o tradutor de texto
subscripiton_key = 'YOUR_SUBSCRIPTION_KEY'
endpoint = 'https://api.cognitive.microsofttranslator.com'
location = 'YOUR_LOCATION'
target_language = 'pt-br'

def translate_text(text,target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscripiton_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom())
    }
    body = [{
        'text': text
    }]
    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': target_language
    }
    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    return response[0]["translations"][0]["text"]


# Traduzindo o documento

    
def translate_document(path):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:
        translated_text = translate_text(paragraph.text, target_language)
        full_text.append(translated_text)

    document = Document()
    for line in full_text:
        translate_document.add_paragraph(line)
    path_translated = path.replace('.docx', f'_{target_language}.docx')

    translate_document.save(path_translated)
    return path_translated



# Extraindo de uma UML



def extract_text_from_url(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        for script_or_style in soup(['script', 'style']):
            script_or_style.decompose()
        texto = soup.get_text()
        linhas = (line.strip() for line in texto.splitlines())
        chunks = (phrase.strip() for line in linhas for phrase in line.split("  "))
        texto_limpo = '\n'.join(chunk for chunk in chunks if chunk)
        return texto_limpo
    else:
        raise ValueError(f"Failed to fetch URL: {url}")



# Utilizando langchain

client = AzureChatOpenAI(
    azure_endpoint = 'YOUR_ENDPOINT',
    api_key = 'YOUR_API_KEY',
    api_version = 'YOUR_API_VERSION',
    deployment_name = 'YOUR_DEPLOYMENT_NAME',
    max_retries=0
)

def translate_article(text,lang):
   messages = [
       ("system","Você atua como um tradutor de texto"),
        ("user",f"traduza o {text}  para o idioma {lang} e responda em markdown"),

   ]

   response = client.invoke(messages)
   return response.content



''' Aqui dara um erro esperado pois tem que ser criado as credenciais no azure'''


url ='Passe a url aqui'
text = extract_text_from_url(url)
article = translate_article(text,'escolha a linguagem')
print(article)